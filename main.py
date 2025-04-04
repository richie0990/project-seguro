import os.path
import sys 
from PyQt6.QtWidgets import  QApplication,QWidget,QLabel,QVBoxLayout,QMessageBox,QMenuBar
from PyQt6.QtGui import QAction,QPixmap,QIcon
from PyQt6.QtCore import Qt
from docx import Document
import pandas as pd
import os
from inicio import Inicio
from agregar import Agregar
from eliminar import Eliminar
import re
from actualizar import Actualizar
from openpyxl import load_workbook

#crear ruta si no existe 
escritorio= os.path.join(os.path.expanduser("~"),"Desktop")
if os.path.exists(escritorio):
    escritorio= os.path.join(os.path.expanduser("~"),"Desktop")
else:
    escritorio= os.path.join(os.path.expanduser("~"),"Escritorio")
path =os.path.isdir(os.path.join(escritorio,"Seguros Atlanta"))
if path :
    ""
else:
    os.mkdir(os.path.join(escritorio,"Seguros Atlanta"))
    

doc = Document(os.path.join(escritorio,"Seguros Atlanta/Seguros.docx"))
path = os.path.join(escritorio,"Seguros Atlanta""/HOJA DE SEGURO ATLANTICA.xlsx")
data_execl = pd.read_excel(path,sheet_name="Sheet1")
heads = list(data_execl)
data = []
j=0

filas, columnas = data_execl.shape
for i in range(filas):
    data.append({})

for sheet_name, df in data_execl.items():   
    for i,valor in enumerate(list(df)):
        
        if not re.match(r"^nan$",str(valor)):
            if re.match(r"^FECHA",str(sheet_name)):
                date = pd.to_datetime(valor,dayfirst=True)
                time = date.strftime('%d/%m/%Y')
                valor= time       
        data[i][heads[j]] = valor
    if j>= len(heads):
        j=0
    else:
        j+=1

class Ventana(QWidget):
    def __init__(self):
        super().__init__()
        self.width =600
        self.height =580
        self.setGeometry(int(self.height-(self.height/2)),int(self.width-(self.width/2)),self.width,self.height)
        self.setWindowIcon(QIcon("./src/img/favicon.ico"))
        self.setWindowTitle("Seguros Atlanta")
        self.heads = heads
        self.valores = data
        self.start_window = Inicio(self)
        self.eliminar_window =  Eliminar(self)
        self.agregar_window = Agregar(self)
        self.agregar_window = Actualizar(self)
        self.re_fechas =    r'^\d{1,2}/\d{1,2}/\d{4}$'
        self.re_cedula =    r'^\d{3}-\d{7}-\d{1}$'
        self.re_telefono =  r'^\d{3}-\d{3}-\d{4}$'
        self.option={
            "current":"inicio",
            "ventana":self.start_window
        } 

        self.msj = QMessageBox()
        self.menu_bar = QMenuBar(self)
        
       #pixmap
        pixmap = QPixmap("./src/img/logo.png")

        #layaout
        self.layout_ = QVBoxLayout()
        
        #titulo
        titulo_width  = 660
        titulo_height =150
        titulo = QLabel("",self)
        titulo.setAlignment(Qt.AlignmentFlag.AlignCenter)
        titulo.setGeometry(int(self.width/6),30,titulo_width,titulo_height)
        scale_pixmap= pixmap.scaled(titulo_width,titulo_height,Qt.AspectRatioMode.KeepAspectRatio)
        titulo.setPixmap(scale_pixmap)

        titulo.setFixedHeight(titulo_height)

        #add layout
        self.layout_.addWidget(titulo)
        self.layout_.addWidget(self.start_window)
        self.setLayout(self.layout_)

        #setmenuebar
        self.layout_.setMenuBar(self.menu_bar)
        
        #crear bar menu
        file_menu = self.menu_bar.addMenu("Archivo")

        #agregar acctiones
        agregar_action = QAction("Agregar",self)
        eliminar_action = QAction("Eliminar",self)
        salir_action = QAction("Salir",self)
        inicio_action = QAction("Inicio",self)
        actualizar_action = QAction("Actualizar",self)

        # conectando los acctiones
        inicio_action.triggered.connect(self.inicio)
        agregar_action.triggered.connect(self.agregar)
        eliminar_action.triggered.connect(self.eliminar)
        salir_action.triggered.connect(self.salir) 
        actualizar_action.triggered.connect(self.actualizar)       

        #agregarla acction al menu
        file_menu.addAction(inicio_action)
        file_menu.addAction(agregar_action)
        file_menu.addAction(eliminar_action)
        file_menu.addAction(actualizar_action)
        file_menu.addAction(salir_action)
        
    def salir(self):
        sys.exit()
    
    def agregar(self):

        self.option["current"] = "agregar"
        self.crear_windows()
        self.layout_.addWidget(self.agregar_window)

    def eliminar(self):
       
        self.option["current"] = "eliminar"
        self.crear_windows() 
        self.layout_.addWidget(self.eliminar_window)

    def inicio(self):
        self.option["current"] = "inicio"
        self.crear_windows() 
        self.layout_.addWidget(self.start_window)

    def actualizar(self):
        self.option["current"] = "actualizar"
        self.crear_windows() 
        self.layout_.addWidget(self.actualizar_window)

    def crear_windows(self):
        if self.option["current"] == "eliminar" :
            self.layout_.removeWidget(self.option["ventana"])
            self.option["ventana"].deleteLater()  
            self.eliminar_window= Eliminar(self)
            self.option["ventana"]=self.eliminar_window

        elif self.option["current"] == "inicio" :
            self.layout_.removeWidget(self.option["ventana"])
            self.option["ventana"].deleteLater()  
            self.start_window = Inicio(self)
            self.option["ventana"]=self.start_window

        elif self.option["current"] == "agregar" :
            self.layout_.removeWidget(self.option["ventana"])
            self.option["ventana"].deleteLater()  
            self.agregar_window= Agregar(self)
            self.option["ventana"]=self.agregar_window
        elif self.option["current"] == "actualizar" :
            self.layout_.removeWidget(self.option["ventana"])
            self.option["ventana"].deleteLater()  
            self.actualizar_window= Actualizar(self)
            self.option["ventana"]=self.actualizar_window
        
        
    def save_word(self,data):
        
        archivo =   os.path.join(escritorio,"Seguros Atlanta/Seguros.docx")
        tables = doc.tables
        if tables:
            table_to_remove = tables[0]
            tbl = table_to_remove._element  # Accedemos al elemento XML de la tabla
            tbl.getparent().remove(tbl) 
        tabla = doc.add_table(rows=len(data),cols=7)

        tabla.style = 'Table Grid'
        for i, fila in enumerate(data):
            for j, valor in enumerate(fila):
                # Acceder a cada celda y asignar el valor
                tabla.cell(i, j).text = valor
        # Guardar el documento
       
        doc.save(archivo)
        #solo hay que poner donde el quiere que se guarde la informacion 
    def save_excel(self,data):
        archivo =   os.path.join(escritorio,"Seguros Atlanta/HOJA DE SEGURO ATLANTICA.xlsx")
        #print(data)
        df = pd.DataFrame(data)
        df.to_excel(archivo,index=False)
        # Cargar el archivo Excel con openpyxl
        wb = load_workbook(archivo)
        ws = wb.active

        # Ajustar el tamaño de las columnas para que se ajusten al contenido de las celdas
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter  # Obtener la letra de la columna

            for cell in col:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(cell.value)  # Encontrar el valor más largo
                except:
                    pass
                
            adjusted_width = (max_length + 2)  # Ajustar un poco el tamaño (puedes cambiar el número)
            ws.column_dimensions[column].width = adjusted_width

        # Guardar el archivo de Excel con los tamaños ajustados
        wb.save(archivo)
    def convert_to_excel(self):
        converted_data  = {}
        for head in self.heads:
            converted_data[head] = []
        #####AQui me quede hay que probar los datos de que s epuedan convertir en execl data.

        for i,valor_obj in enumerate(self.valores):
            
            for j,valor in enumerate(list(valor_obj.values())):
                
                converted_data[self.heads[j]].append(valor)
        return converted_data 
                 
if __name__ == "__main__":
    
    app = QApplication(sys.argv)
    mi_ventana= Ventana()
    mi_ventana.showFullScreen()
    sys.exit(app.exec())


    